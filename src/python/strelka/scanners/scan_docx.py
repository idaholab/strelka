import io
import zipfile

from bs4 import BeautifulSoup
import docx

from . import File, Options, Scanner
from ..model import Date


class ScanDocx(Scanner):
    """Collects metadata and extracts text from docx files.

    Options:
        extract_text: Boolean that determines if document text should be
            extracted as a child file.
            Defaults to False.
    """

    def scan(self, data: bytes, file: File, options: Options, expire_at: Date) -> None:
        extract_text = options.get("extract_text", False)

        with io.BytesIO(data) as docx_io:
            try:
                doc = docx.Document(docx_io)
                self.event.update(
                    {
                        "total": {
                            "words": 0,
                            "images": 0,
                        },
                        "author": doc.core_properties.author,
                        "category": doc.core_properties.category,
                        "comments": doc.core_properties.comments,
                        "content_status": doc.core_properties.content_status,
                        "created": doc.core_properties.created,
                        "identifier": doc.core_properties.identifier,
                        "keywords": doc.core_properties.keywords,
                        "language": doc.core_properties.language,
                        "last_modified_by": doc.core_properties.last_modified_by,
                        "last_printed": doc.core_properties.last_printed,
                        "modified": doc.core_properties.modified,
                        "revision": doc.core_properties.revision,
                        "subject": doc.core_properties.subject,
                        "title": doc.core_properties.title,
                        "version": doc.core_properties.version,
                        "font_colors": set(),
                    }
                )

                for paragraph in doc.paragraphs:
                    soup = BeautifulSoup(paragraph.paragraph_format.element.xml, "xml")
                    self.event["font_colors"].update(
                        c.attrs["w:val"] for c in soup.select("color")
                    )
                    self.event["total"]["images"] += sum(
                        1 for i in soup.select("pic") if i.attrs["xmlns:pic"]
                    )
                    if paragraph.text.strip():
                        self.event["total"]["words"] += len(paragraph.text.split(" "))

                self.event["white_text_in_doc"] = "FFFFFF" in self.event["font_colors"]

                # send document text contents back to Strelka if requested
                if extract_text:
                    self.emit_file(
                        "\n".join(p.text for p in doc.paragraphs).encode("utf-8"),
                        name=":docx-text-contents",
                    )

            except zipfile.BadZipFile:
                self.add_flag("bad_zip")
            except Exception:
                self.add_flag("bad_doc")
